# ==============================================================================
# CISCO VIRTUAL INTERNSHIP 2026 - CYBER SECURITY PROJECT
# Automated Report Generator: Python to Microsoft Word (.docx)
# Generates: Master Technical Report & Student Contribution Summary
# ==============================================================================

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_color):
    """Set background color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def generate_master_report():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("CISCO VIRTUAL INTERNSHIP 2026\nCYBER SECURITY CAPSTONE PROJECT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x50, 0x73)

    sub_title = doc.add_paragraph()
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_title.add_run("Secure Hybrid Data Center Network Architecture & Multi-Cloud Workload Protection")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Program & Body", "Cisco Virtual Internship 2026 / AICTE & Cisco Networking Academy"),
        ("Problem Statement", "Cyber Security: Enterprise Hybrid Data Center Security & Workload Protection"),
        ("Core Technologies", "Cisco ASA/Firepower, Cisco ISE, Cisco Duo, AWS TGW, Calico CNI, Istio mTLS"),
        ("Submission Deadline", "30th August 2026"),
        ("Document Type", "Technical Architecture Specification & Capstone Report")
    ]
    for i, (k, v) in enumerate(meta_data):
        cell_0 = meta_table.cell(i, 0)
        cell_1 = meta_table.cell(i, 1)
        cell_0.text = k
        cell_1.text = v
        set_cell_background(cell_0, "EBF3F9")
        set_cell_background(cell_1, "F8FAFC")
        cell_0.paragraphs[0].runs[0].font.bold = True
        cell_0.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell_1.paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph()

    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "Modern enterprise and higher education institutions are transitioning from isolated on-premises "
        "data centers to distributed hybrid multi-cloud architectures. Today's academic and enterprise "
        "workloads span private data centers (Red Hat OpenShift, local database clusters) and public cloud "
        "environments (AWS EKS, RDS, Azure AKS). While this hybrid paradigm unlocks unprecedented scalability "
        "and collaborative power, it fundamentally dissolves the traditional network perimeter.\n\n"
        "This project presents an enterprise-grade Secure Hybrid Data Center Network Architecture designed to "
        "balance simplicity, rock-solid security, and linear scalability. Leveraging Cisco enterprise security "
        "solutions (Cisco ASA 5506-X / Firepower, Cisco ISE, Cisco Duo MFA, Cisco Stealthwatch) combined with "
        "cloud-native controls (AWS Transit Gateway, SG-to-SG referencing, Calico CNI default-deny, and Istio mTLS), "
        "this architecture guarantees that any breach within a single container, virtual machine, or remote endpoint "
        "is immediately contained, logged, and isolated."
    )

    h1 = doc.add_heading("2. Problem Statement & Key Objectives", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "The project addresses the following critical enterprise cybersecurity challenges:\n"
        "• Hybrid Interconnect Protection: Securing data traversing private data centers and public cloud infrastructure via IKEv2 IPsec VPN.\n"
        "• Unified Identity & Access Management (IAM): Eliminating static credentials through OIDC Federation (EKS IRSA), Least-Privilege RBAC, and Cisco Duo MFA.\n"
        "• Multi-VPC Micro-Segmentation: Designing a Hub-and-Spoke Transit Gateway network with 3-tier subnets and SG-to-SG referencing.\n"
        "• Kubernetes & Container Hardening: Enforcing Calico CNI default-deny policies, Restricted Pod Security Standards, and Istio STRICT mTLS.\n"
        "• Secure Remote Faculty Access: Deploying Cisco AnyConnect ZTNA with dynamic Cisco ISE posture assessments.\n"
        "• DevSecOps Governance: Establishing automated CI/CD security gates (SAST, Checkov IaC, Trivy CVE scanning, Cosign signing)."
    )

    h1 = doc.add_heading("3. Network Topology & IP Addressing Scheme", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)

    ip_table = doc.add_table(rows=1, cols=4)
    ip_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = ip_table.rows[0].cells
    hdr_titles = ["Segment Name", "Subnet CIDR", "VLAN / Zone", "Security Function"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "005073")
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)

    ip_rows = [
        ("Enterprise Core Transit", "10.10.0.0/30", "P2P", "Core Router to ASA Inside link"),
        ("Faculty Workstations", "10.10.10.0/24", "VLAN 10", "On-Premises Faculty Desktops (802.1X)"),
        ("Student & Staff LAN", "10.10.20.0/24", "VLAN 20", "General Campus traffic; DB probing restricted"),
        ("On-Premises DC / OpenShift", "10.10.30.0/24", "VLAN 30", "Private Data Center servers & OpenShift"),
        ("Network Management", "10.10.99.0/24", "VLAN 99", "Out-of-band admin jumpboxes & SSH"),
        ("Faculty AnyConnect Pool", "172.16.50.0/24", "SSL VPN", "Dynamic IP pool for remote ZTNA faculty"),
        ("Enterprise Edge Outside", "203.0.113.0/30", "Outside (0)", "Cisco ASA Public IP (203.0.113.2)"),
        ("Cloud Transit WAN Edge", "198.51.100.0/30", "Cloud WAN", "Cloud Transit Gateway (198.51.100.2)"),
        ("Cloud Tier 1 (ALB / Ingress)", "172.20.1.0/24", "Public Subnet", "AWS Application Load Balancers & WAF"),
        ("Cloud Tier 2 (EKS Microservices)", "172.20.2.0/24", "Private Subnet", "Container worker nodes; zero public IP"),
        ("Cloud Tier 3 (Isolated DB)", "172.20.3.0/24", "Isolated Subnet", "Managed DB (RDS); zero internet routing")
    ]
    for r in ip_rows:
        row_cells = ip_table.add_row().cells
        for i, val in enumerate(r):
            row_cells[i].text = val
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_background(row_cells[i], "F8FAFC" if len(ip_table.rows) % 2 == 0 else "FFFFFF")

    doc.add_paragraph()

    h1 = doc.add_heading("4. Cloud Security Groups & IAM Architecture", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "To mitigate lateral attack propagation across cloud workloads, the architecture implements two foundational controls:\n"
        "1. SG-to-SG Referencing: Security groups reference peer Security Group IDs rather than CIDR blocks. Inbound traffic to the EKS Microservices tier is accepted ONLY if originating from the ALB Security Group (sg-prod-ingress-alb) on port 8080. Database ingress is strictly permitted ONLY from sg-prod-eks-microservices on port 5432.\n"
        "2. EKS IAM Roles for Service Accounts (IRSA): Microservice pods authenticate to AWS services using temporary, cryptographic STS tokens issued via OIDC federation. No static access keys exist inside containers, neutralizing credential theft risk."
    )

    h1 = doc.add_heading("5. Container & Kubernetes Security (OpenShift / EKS)", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "• Calico CNI Default-Deny: A baseline NetworkPolicy drops all ingress and egress traffic by default in the production namespace. Explicit rules permit only frontend-to-backend and backend-to-database communication.\n"
        "• Pod Security Standards (PSS): Enforces the 'Restricted' profile requiring containers to run as non-root (UID 10001), with read-only root filesystems, dropped Linux capabilities (drop ALL), and default seccomp profiles.\n"
        "• Istio Service Mesh mTLS: Enforces STRICT mutual TLS 1.3 encryption for all pod-to-pod streams with SPIFFE cryptographic X.509 identity validation and method-level AuthorizationPolicy rules."
    )

    h1 = doc.add_heading("6. Secure Remote Faculty Access (ZTNA)", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "Faculty members connecting from home or campus utilize Cisco AnyConnect / Cisco Secure Client with Zero Trust Network Access:\n"
        "• Cisco Duo MFA: Mandatory out-of-band push verification or FIDO2 hardware token.\n"
        "• Cisco ISE Posture Validation: Verifies OS patch level, BitLocker/FileVault disk encryption, and active Cisco Secure Endpoint EDR prior to granting network access.\n"
        "• Optimized Split-Tunneling: Only enterprise campus and cloud VPC subnets are routed across the encrypted tunnel, preserving institutional bandwidth for video conferencing and personal traffic."
    )

    h1 = doc.add_heading("7. Multi-Stakeholder Collaboration & DevSecOps", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "A clear RACI governance model coordinates Application Developers, Network Designers, Kubernetes Platform Engineers, and Security Teams. "
        "The automated CI/CD pipeline enforces 4 strict security gates (Semgrep SAST, Checkov IaC scanning, Trivy container CVE scanning, and Cosign image signing) "
        "before ArgoCD GitOps synchronizes code to production."
    )

    h1 = doc.add_heading("8. Architectural Evaluation: Simplicity, Security & Scale", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "• Simplicity: Transit Gateway hub eliminates full-mesh complexity; declarative GitOps and unified SSO streamline operations.\n"
        "• Security: Defense-in-depth across physical, virtual, and container layers with zero static credentials and micro-segmented blast radius.\n"
        "• Scale: Linear scalability supporting thousands of microservice pods and up to 50 Gbps throughput per VPC attachment without redesign."
    )

    output_path = r"d:\PROJECTS\CISCO\docs\Cisco_Cybersecurity_Internship_Report.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

def generate_summary_document():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("CISCO VIRTUAL INTERNSHIP 2026\nINDIVIDUAL PROJECT CONTRIBUTION SUMMARY")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x00, 0x50, 0x73)

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Candidate Name", "[Your Full Name]"),
        ("AICTE Registration ID", "[Your AICTE ID, e.g., STUXXXXXXXXXX]"),
        ("College Name", "[Your College Name, e.g., GL Bajaj Institute of Technology & Management]"),
        ("Technology Track", "Cyber Security & Enterprise Networking"),
        ("Project Title", "Secure Hybrid Data Center Network Architecture & Multi-Cloud Workload Protection")
    ]
    for i, (k, v) in enumerate(meta_data):
        cell_0 = meta_table.cell(i, 0)
        cell_1 = meta_table.cell(i, 1)
        cell_0.text = k
        cell_1.text = v
        set_cell_background(cell_0, "EBF3F9")
        set_cell_background(cell_1, "F8FAFC")
        cell_0.paragraphs[0].runs[0].font.bold = True
        cell_0.paragraphs[0].runs[0].font.size = Pt(9.5)
        cell_1.paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph()

    h1 = doc.add_heading("1. Individual Contribution Breakdown", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "• Hybrid IPsec VPN & Edge Firewall: Configured Site-to-Site IKEv2 IPsec VPN (AES-256-GCM / SHA-256 / DH 14) and Identity NAT on Cisco ASA 5506-X to seamlessly connect on-premise OpenShift clusters to AWS Transit Gateway.\n"
        "• Cloud VPC Segmentation & Security Groups: Authored Terraform code for 3-tier subnets and SG-to-SG referencing rules, ensuring complete isolation of database tiers from public ingress.\n"
        "• Kubernetes Microservice Security: Created Calico default-deny NetworkPolicies, Restricted Pod Security Standards, and Istio STRICT mTLS 1.3 manifests.\n"
        "• Zero-Trust Remote Faculty Access: Designed Cisco AnyConnect SSL VPN profiles with Cisco Duo MFA push notifications and Cisco ISE posture compliance assessment.\n"
        "• DevSecOps Pipeline Automation: Integrated Semgrep SAST, Checkov IaC scanning, and Trivy container scanning into automated CI/CD security gates.\n"
        "• Packet Tracer Lab Simulation: Built and validated the entire multi-device network topology with complete CLI configuration scripts."
    )

    h1 = doc.add_heading("2. Cisco Packet Tracer Deliverables", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x00, 0x50, 0x73)
    p = doc.add_paragraph(
        "The submission includes tested configuration files for Enterprise Core Router, ASA Perimeter Firewall, "
        "Data Center Switch, ISP Router, Cloud Transit Gateway Router, and Cloud Workload Switch."
    )

    doc.add_paragraph("\nDeclaration: I hereby declare that this project work is original and completed in accordance with Cisco Networking Academy and AICTE guidelines.\n")
    doc.add_paragraph("Candidate Signature: _______________________          Date: 24th August 2026")

    output_path = r"d:\PROJECTS\CISCO\docs\Student_Contribution_Summary.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    generate_master_report()
    generate_summary_document()
